from flask import Flask, render_template, request, redirect, session, send_from_directory, flash, url_for, jsonify, send_file
import os
from werkzeug.utils import secure_filename
from datetime import datetime, timedelta
import json
from werkzeug.security import check_password_hash, generate_password_hash
import random
from docx import Document
import io
from bs4 import BeautifulSoup
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from base64 import b64encode
import re, base64
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from openpyxl import Workbook
import requests
import uuid

app = Flask(__name__)
app.secret_key = 'secret-key'  # ganti di produksi
UPLOAD_FOLDER = 'uploads'
MAX_STORAGE = 5 * 1024 * 1024 * 1024  # 5GB
RECYCLE_BIN = 'Recycle_Bin'
MAINTENANCE_MODE = False  # Set True jika ingin aktifkan mode maintenance
app.permanent_session_lifetime = timedelta(hours=24)

users = {}  # Format: {email: hashed_password}
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

if not os.path.exists(UPLOAD_FOLDER):
    os.makedirs(UPLOAD_FOLDER, exist_ok=True)

USER_FILE = 'users.json'

DEVELOPER_EMAIL = 'cloudfiledev104421@mail.com'
DEVELOPER_PASSWORD = 'hellnowdev8119001'

def load_users():
    global users
    if os.path.exists(USER_FILE):
        try:
            with open(USER_FILE, 'r') as f:
                content = f.read().strip()
                if content:
                    users = json.loads(content)
                else:
                    users = {}
        except Exception as e:
            print(f"Error loading users.json: {e}")
            users = {}
    else:
        users = {}

def save_users():
    with open(USER_FILE, 'w') as f:
        json.dump(users, f)

# Panggil load_users() saat aplikasi mulai
load_users()

@app.context_processor
def inject_user():
    email = session.get('user')
    user_name = None
    # Jika developer, ambil nama dari session jika ada
    if session.get('is_developer'):
        user_name = session.get('user_name') or 'Developer'
        return dict(email=email, user_name=user_name)
    if email:
        user = users.get(email)
        if isinstance(user, dict):
            user_name = user.get('name')
    return dict(email=email, user_name=user_name)

def get_user_folder(email):
    folder = os.path.join(UPLOAD_FOLDER, email)
    if os.path.exists(folder):
        if not os.path.isdir(folder):
            raise Exception(f"Path {folder} exists and is not a directory!")
    else:
        try:
            os.makedirs(folder, exist_ok=True)
        except FileExistsError:
            if not os.path.isdir(folder):
                raise
    return folder

def get_storage_used(folder):
    total_size = 0
    for dirpath, dirnames, filenames in os.walk(folder):
        for f in filenames:
            fp = os.path.join(dirpath, f)
            try:
                if os.path.isfile(fp):
                    total_size += os.path.getsize(fp)
            except Exception as e:
                # Bisa log error jika perlu
                pass
    return total_size

def get_recycle_bin_folder(email):
    folder = os.path.join(UPLOAD_FOLDER, email, RECYCLE_BIN)
    if os.path.exists(folder):
        if not os.path.isdir(folder):
            raise Exception(f"Path {folder} exists and is not a directory!")
    else:
        try:
            os.makedirs(folder, exist_ok=True)
        except FileExistsError:
            if not os.path.isdir(folder):
                raise
    return folder

@app.before_request
def check_maintenance():
    if MAINTENANCE_MODE and request.endpoint != 'maintenance':
        return redirect('/maintenance')

@app.before_request
def check_session_token():
    if 'user' in session:
        user = users.get(session['user'])
        if user and 'session_token' in user:
            if session.get('session_token') != user['session_token']:
                session.clear()
                flash('Anda telah logout dari perangkat lain.')
                return redirect('/login')

@app.route('/')
def index():
    if 'user' in session:
        return redirect('/dashboard')
    return render_template('homepage.html')

@app.route('/register', methods=['GET', 'POST'])
def register():
    if request.method == 'POST':
        name = request.form['name'].strip()
        email = request.form['email'].strip().lower()
        password = request.form['password'].strip()
        confirm_password = request.form['confirm_password'].strip()
        q1 = request.form['q1']
        a1 = request.form['a1'].strip()
        q2 = request.form['q2']
        a2 = request.form['a2'].strip()
        q3 = request.form['q3']
        a3 = request.form['a3'].strip()
        if email in users:
            flash('Email sudah terdaftar. Silakan gunakan email lain atau login.')
            return render_template('register.html', email=email, name=name)
        if len(password) < 8:
            flash('Password minimal 8 karakter.')
            return render_template('register.html', email=email, name=name)
        if password != confirm_password:
            flash('Password dan konfirmasi password tidak sama.')
            return render_template('register.html', email=email, name=name)
        users[email] = {
            'name': name,
            'password': generate_password_hash(password),
            'q1': q1,
            'a1': generate_password_hash(a1),
            'q2': q2,
            'a2': generate_password_hash(a2),
            'q3': q3,
            'a3': generate_password_hash(a3),
            'package': request.form.get('package', 'free'),
            'status': 'Plus S' if request.form.get('package') == 'S' else 'Free'
        }
        save_users()
        get_user_folder(email)
        flash('Registration successful. Please login.')
        return redirect('/login')
    return render_template('register.html')

@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        turnstile_token = request.form.get('cf-turnstile-response')
        if not turnstile_token:
            flash('Captcha wajib diisi.')
            return render_template('login.html')
        secret_key = '0x4AAAAAABjKlC5xKpUVPv2LDubSGrsc2FI'  # Ganti dengan secret key asli Anda
        verify_resp = requests.post('https://challenges.cloudflare.com/turnstile/v0/siteverify', data={
            'secret': secret_key,
            'response': turnstile_token,
            'remoteip': request.remote_addr
        })
        verify_result = verify_resp.json()
        print('Turnstile verify result:', verify_result)
        if not verify_result.get('success'):
            flash('Verifikasi captcha gagal. Silakan coba lagi.')
            return render_template('login.html')
        email = request.form['email'].strip().lower()
        password = request.form['password'].strip()
        # Akun pengembang
        if email == DEVELOPER_EMAIL and password == DEVELOPER_PASSWORD:
            session['user'] = email
            session['is_developer'] = True
            session['session_token'] = str(uuid.uuid4())
            session.permanent = True
            flash('Berhasil login.')
            return redirect('/dashboard')
        stored_user = users.get(email)
        print(f"Login attempt: {email} / {password}")
        print(f"Stored user: {stored_user}")
        if not stored_user or not check_password_hash(stored_user['password'], password):
            flash('Email atau password salah.')
            return render_template('login.html')
        session_token = str(uuid.uuid4())
        session['user'] = email
        session['session_token'] = session_token
        users[email]['session_token'] = session_token
        save_users()
        session.permanent = True  # Aktifkan session permanent (timeout 24 jam)
        # Simpan info perangkat
        device_id = str(uuid.uuid4())
        user_agent = request.headers.get('User-Agent', 'Unknown')
        ip = request.remote_addr
        login_time = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        device_info = {
            'device_id': device_id,
            'user_agent': user_agent,
            'ip': ip,
            'login_time': login_time,
            'session_token': session_token
        }
        if 'devices' not in users[email]:
            users[email]['devices'] = []
        # Batasi maksimal 3 perangkat
        if len(users[email]['devices']) >= 3:
            users[email]['devices'] = users[email]['devices'][-2:]  # Sisakan 2 terbaru
        users[email]['devices'].append(device_info)
        save_users()
        session['device_id'] = device_id
        if email == 'cloudfiledev104421@mail.com':
            session['user_has_package'] = True
            session['user_package_code'] = 'M'
        flash('Berhasil login.')
        return redirect('/dashboard')
    return render_template('login.html')

@app.route('/logout')
def logout():
    session.pop('user', None)
    return redirect('/login')

@app.route('/logout_all', methods=['POST'])
def logout_all():
    if 'user' in session:
        email = session['user']
        if email in users:
            users[email]['session_token'] = str(uuid.uuid4())  # Reset token
            save_users()
        session.clear()
    return redirect('/login')

def get_storage_limit_by_package(package_code):
    if package_code == 'M':
        return 100 * 1024 * 1024 * 1024  # 100GB
    elif package_code == 'S':
        return 50 * 1024 * 1024 * 1024   # 50GB
    elif package_code == 'L':
        return 200 * 1024 * 1024 * 1024  # 200GB (contoh)
    return 5 * 1024 * 1024 * 1024        # Default 5GB

@app.route('/dashboard', defaults={'path': ''}, methods=['GET', 'POST'])
@app.route('/dashboard/<path:path>', methods=['GET', 'POST'])
def dashboard(path):
    if 'user' not in session:
        return redirect('/login')
    email = session['user']
    if email == 'cloudfiledev104421@mail.com':
        session['user_has_package'] = True
        session['user_package_code'] = 'M'
    else:
        session['user_has_package'] = False
        session['user_package_code'] = 'S'
    user_folder = get_user_folder(email)
    current_folder = os.path.join(user_folder, path)
    if os.path.exists(current_folder):
        if not os.path.isdir(current_folder):
            raise Exception(f"Path {current_folder} exists and is not a directory!")
    else:
        try:
            os.makedirs(current_folder, exist_ok=True)
        except FileExistsError:
            if not os.path.isdir(current_folder):
                raise

    # Upload file ke folder saat ini
    if request.method == 'POST':
        if 'file' in request.files:
            file = request.files['file']
            if file and file.filename:
                filename = secure_filename(file.filename or "")
                file_path = os.path.join(current_folder, filename)
                current_size = get_storage_used(user_folder)
                file.seek(0, os.SEEK_END)
                file_size = file.tell()
                file.seek(0)
                if current_size + file_size > MAX_STORAGE:
                    flash('Storage limit exceeded (5GB)')
                    return redirect(request.url)
                file.save(file_path)
                flash('File uploaded')
        elif 'foldername' in request.form:
            foldername = secure_filename(request.form['foldername'])
            folder_path = os.path.join(current_folder, foldername)
            if os.path.exists(folder_path):
                if not os.path.isdir(folder_path):
                    raise Exception(f"Path {folder_path} exists and is not a directory!")
            else:
                try:
                    os.makedirs(folder_path, exist_ok=True)
                except FileExistsError:
                    if not os.path.isdir(folder_path):
                        raise
            flash('Folder created')
        return redirect(request.url)

    # List isi folder
    items = os.listdir(current_folder)
    files = []
    folders = []
    for item in items:
        if item == RECYCLE_BIN:
            continue  # Sembunyikan folder Recycle_Bin
        item_path = os.path.join(current_folder, item)
        if os.path.isdir(item_path):
            folders.append(item)
        else:
            files.append(item)

    # --- FILTER FILES BY SEARCH ---
    q = request.args.get('q', '').lower()
    if q:
        files = [f for f in files if q in f.lower()]

    # --- SORT FILES BY TIME ---
    sort = request.args.get('sort', '')
    if sort == 'newest':
        files.sort(key=lambda f: os.path.getmtime(os.path.join(current_folder, f)), reverse=True)
    elif sort == 'oldest':
        files.sort(key=lambda f: os.path.getmtime(os.path.join(current_folder, f)))

    files_ok, sizes, times = [], [], []
    for f in files:
        try:
            fp = os.path.join(current_folder, f)
            sizes.append(os.path.getsize(fp))
            times.append(datetime.fromtimestamp(os.path.getmtime(fp)).strftime('%Y-%m-%d %H:%M'))
            files_ok.append(f)
        except Exception:
            continue

    total_used = get_storage_used(user_folder)
    view = request.args.get('view', 'grid')  # default grid
    user_has_package = session.get('user_has_package', False)
    user_package_code = session.get('user_package_code', 'S' if not user_has_package else session.get('user_package_code', 'M'))
    storage_limit = get_storage_limit_by_package(user_package_code)
    return render_template('dashboard.html', files=list(zip(files_ok, sizes, times)), folders=folders, used=total_used, path=path, email=email, view=view, user_has_package=user_has_package, user_package_code=user_package_code, storage_limit=storage_limit)

@app.route('/download/<path:path>')
def download(path):
    if 'user' not in session:
        return redirect('/login')
    return send_from_directory(get_user_folder(session['user']), path, as_attachment=True)

@app.route('/delete/<filename>')
def delete(filename):
    if 'user' not in session:
        return redirect('/login')
    src_path = os.path.join(get_user_folder(session['user']), filename)
    recycle_folder = get_recycle_bin_folder(session['user'])
    if os.path.exists(src_path):
        base, ext = os.path.splitext(os.path.basename(filename))
        dst_path = os.path.join(recycle_folder, os.path.basename(filename))
        i = 1
        while os.path.exists(dst_path):
            dst_path = os.path.join(recycle_folder, f"{base}({i}){ext}")
            i += 1
        os.rename(src_path, dst_path)
        flash('File dipindahkan ke Recycle Bin')
    return redirect('/dashboard')

@app.route('/profile', methods=['POST'])
def profile_post():
    if 'user' not in session:
        return jsonify({'success': False, 'message': 'Not logged in'}), 401
    email = session['user']
    old_password = request.form.get('old_password') or ''
    new_password = request.form.get('new_password') or ''
    stored_user = users.get(email)
    if not (isinstance(stored_user, dict) and check_password_hash(stored_user['password'], old_password)):
        return jsonify({'success': False, 'message': 'Old password is incorrect'})
    users[email]['password'] = generate_password_hash(new_password)
    save_users()
    return jsonify({'success': True, 'message': 'Password updated successfully'})

@app.route('/preview/<filename>')
def preview(filename):
    if 'user' not in session:
        return redirect('/login')
    user_folder = get_user_folder(session['user'])
    file_path = os.path.join(user_folder, filename)
    if not os.path.exists(file_path):
        flash('File not found')
        return redirect('/dashboard')
    ext = os.path.splitext(filename)[1].lower()
    docx_html = None
    excel_html = None
    if ext == '.docx':
        try:
            from docx import Document
            doc = Document(file_path)
            html = ''
            for para in doc.paragraphs:
                text = para.text.replace('\n', '<br>')
                style = getattr(para.style, 'name', None)
                style_lower = style.lower() if isinstance(style, str) else ''
                if 'heading' in style_lower:
                    level = ''.join(filter(str.isdigit, style_lower)) or '1'
                    html += f'<h{level}>{text}</h{level}>'
                else:
                    html += f'<p>{text}</p>'
            docx_html = html
        except Exception as e:
            docx_html = f'<div class="text-danger">Gagal membaca file docx: {e}</div>'
    elif ext == '.xlsx':
        try:
            from openpyxl import load_workbook
            wb = load_workbook(file_path, read_only=True)
            ws = getattr(wb, 'active', None)
            if ws is not None:
                max_row = getattr(ws, 'max_row', 0)
                max_col = getattr(ws, 'max_column', 0)
                table = '<div class="table-responsive"><table class="table table-bordered table-sm align-middle bg-white"><thead><tr>'
                # Header kolom A-Z
                for col in range(1, max_col+1):
                    col_letter = chr(64+col) if col <= 26 else f'Col{col}'
                    table += f'<th>{col_letter}</th>'
                table += '</tr></thead><tbody>'
                for row in ws.iter_rows(min_row=1, max_row=max_row, min_col=1, max_col=max_col, values_only=True):
                    table += '<tr>'
                    for cell in row:
                        table += f'<td>{cell if cell is not None else ""}</td>'
                    table += '</tr>'
                table += '</tbody></table></div>'
                excel_html = table
            else:
                excel_html = '<div class="text-danger">Worksheet tidak ditemukan.</div>'
        except Exception as e:
            excel_html = f'<div class="text-danger">Gagal membaca file Excel: {e}</div>'
    return render_template('preview.html', filename=filename, docx_html=docx_html, excel_html=excel_html)

@app.route('/files/<path:filename>')
def serve_file(filename):
    if 'user' not in session:
        return redirect('/login')
    return send_from_directory(get_user_folder(session['user']), filename)

@app.route('/delete_folder/<path:path>')
def delete_folder(path):
    if 'user' not in session:
        return redirect('/login')
    folder_path = os.path.join(get_user_folder(session['user']), path)
    recycle_folder = get_recycle_bin_folder(session['user'])
    if os.path.exists(folder_path) and os.path.isdir(folder_path):
        import shutil
        dst_path = os.path.join(recycle_folder, os.path.basename(path))
        shutil.move(folder_path, dst_path)
        flash('Folder dipindahkan ke Recycle Bin')
    return redirect(url_for('dashboard'))

@app.route('/delete_account', methods=['POST'])
def delete_account():
    if 'user' not in session:
        return jsonify({'success': False, 'message': 'Not logged in'}), 401
    email = session['user']
    # Hapus folder user beserta semua file
    user_folder = get_user_folder(email)
    if os.path.exists(user_folder):
        import shutil
        shutil.rmtree(user_folder)
    # Hapus user dari dict users
    users.pop(email, None)
    save_users()
    session.pop('user', None)
    return jsonify({'success': True, 'message': 'Your account has been deleted.'})

@app.route('/delete_bulk', methods=['POST'])
def delete_bulk():
    if 'user' not in session:
        return jsonify({'success': False}), 401
    data = request.get_json()
    email = session['user']
    user_folder = get_user_folder(email)
    for item in data.get('items', []):
        path = os.path.join(user_folder, item['name'])
        if item['type'] == 'file' and os.path.isfile(path):
            os.remove(path)
        elif item['type'] == 'folder' and os.path.isdir(path):
            import shutil
            shutil.rmtree(path)
    return jsonify({'success': True})

@app.route('/forgot_password', methods=['GET', 'POST'])
def forgot_password():
    step = request.form.get('step', 'email')
    email = request.form.get('email', '').strip().lower()
    user = users.get(email)
    error = None
    if request.method == 'POST':
        if step == 'email':
            if not user:
                error = 'Email tidak ditemukan.'
                return render_template('forgot_password.html', step='email', error=error)
            # Tampilkan pertanyaan
            return render_template('forgot_password.html', step='questions', email=email, q1=user['q1'], q2=user['q2'], q3=user['q3'])
        elif step == 'questions':
            a1 = request.form.get('a1', '').strip()
            a2 = request.form.get('a2', '').strip()
            a3 = request.form.get('a3', '').strip()
            if not (user and check_password_hash(user.get('a1', ''), a1) and check_password_hash(user.get('a2', ''), a2) and check_password_hash(user.get('a3', ''), a3)):
                error = 'Jawaban salah. Silakan coba lagi.'
                return render_template('forgot_password.html', step='questions', email=email, q1=user['q1'], q2=user['q2'], q3=user['q3'], error=error)
            # Tampilkan form ubah password baru
            return render_template('forgot_password.html', step='reset', email=email)
        elif step == 'reset':
            new_password = request.form.get('new_password', '').strip()
            confirm_password = request.form.get('confirm_password', '').strip()
            if len(new_password) < 8:
                error = 'Password minimal 8 karakter.'
                return render_template('forgot_password.html', step='reset', email=email, error=error)
            if new_password != confirm_password:
                error = 'Password dan konfirmasi password tidak sama.'
                return render_template('forgot_password.html', step='reset', email=email, error=error)
            users[email]['password'] = generate_password_hash(new_password)
            save_users()
            flash('Password berhasil diubah. Silakan login dengan password baru.')
            return redirect('/login')
    return render_template('forgot_password.html', step='email')

@app.route('/update_name', methods=['POST'])
def update_name():
    if 'user' not in session:
        return jsonify({'success': False, 'message': 'Not logged in'}), 401
    email = session['user']
    new_name = request.form.get('new_name', '').strip()
    if not new_name:
        return jsonify({'success': False, 'message': 'Nama tidak boleh kosong'})
    # Jika developer, update di session saja
    if session.get('is_developer'):
        session['user_name'] = new_name
        return jsonify({'success': True, 'message': 'Nama berhasil diupdate (developer)', 'new_name': new_name})
    user = users.get(email)
    if not isinstance(user, dict):
        return jsonify({'success': False, 'message': 'User not found'})
    user['name'] = new_name
    save_users()
    return jsonify({'success': True, 'message': 'Nama berhasil diupdate', 'new_name': new_name})

@app.route('/recycle_bin')
def recycle_bin():
    if 'user' not in session:
        return redirect('/login')
    email = session['user']
    recycle_folder = get_recycle_bin_folder(email)
    items = os.listdir(recycle_folder)
    files = []
    folders = []
    for item in items:
        item_path = os.path.join(recycle_folder, item)
        if os.path.isdir(item_path):
            folders.append(item)
        else:
            files.append(item)
    # Info: tampilkan waktu hapus (mtime)
    files_info = [(f, os.path.getsize(os.path.join(recycle_folder, f)), datetime.fromtimestamp(os.path.getmtime(os.path.join(recycle_folder, f))).strftime('%Y-%m-%d %H:%M')) for f in files]
    return render_template('recycle_bin.html', files=files_info, folders=folders, email=email)

@app.route('/restore/<filename>')
def restore_file(filename):
    if 'user' not in session:
        return redirect('/login')
    email = session['user']
    recycle_folder = get_recycle_bin_folder(email)
    file_path = os.path.join(recycle_folder, filename)
    user_folder = get_user_folder(email)
    if os.path.exists(file_path):
        base, ext = os.path.splitext(filename)
        dst_path = os.path.join(user_folder, filename)
        i = 1
        while os.path.exists(dst_path):
            dst_path = os.path.join(user_folder, f"{base}({i}){ext}")
            i += 1
        os.rename(file_path, dst_path)
        flash('File berhasil direstore')
    return redirect('/recycle_bin')

@app.route('/permanent_delete/<filename>')
def permanent_delete(filename):
    if 'user' not in session:
        return redirect('/login')
    email = session['user']
    recycle_folder = get_recycle_bin_folder(email)
    file_path = os.path.join(recycle_folder, filename)
    if os.path.exists(file_path):
        os.remove(file_path)
        flash('File dihapus permanen')
    return redirect('/recycle_bin')

@app.route('/restore_folder/<foldername>')
def restore_folder(foldername):
    if 'user' not in session:
        return redirect('/login')
    email = session['user']
    recycle_folder = get_recycle_bin_folder(email)
    folder_path = os.path.join(recycle_folder, foldername)
    user_folder = get_user_folder(email)
    if os.path.exists(folder_path) and os.path.isdir(folder_path):
        import shutil
        shutil.move(folder_path, os.path.join(user_folder, foldername))
        flash('Folder berhasil direstore')
    return redirect('/recycle_bin')

@app.route('/permanent_delete_folder/<foldername>')
def permanent_delete_folder(foldername):
    if 'user' not in session:
        return redirect('/login')
    email = session['user']
    recycle_folder = get_recycle_bin_folder(email)
    folder_path = os.path.join(recycle_folder, foldername)
    if os.path.exists(folder_path) and os.path.isdir(folder_path):
        import shutil
        shutil.rmtree(folder_path)
        flash('Folder dihapus permanen')
    return redirect('/recycle_bin')

@app.template_filter('filesizeformat')
def filesizeformat(value):
    for unit in ['B', 'KB', 'MB', 'GB', 'TB']:
        if value < 1024.0:
            return f"{value:.1f} {unit}"
        value /= 1024.0
    return f"{value:.1f} PB"

@app.route('/maintenance')
def maintenance():
    return render_template('maintenance.html'), 503

@app.route('/properties', methods=['POST'])
def properties():
    if 'user' not in session:
        return jsonify({'success': False, 'message': 'Not logged in'}), 401
    email = session['user']
    rel_path = (request.json.get('path', '').strip() if request.json else '')
    if '..' in rel_path or rel_path.startswith('/'):
        return jsonify({'success': False, 'message': 'Path tidak valid'}), 400
    abs_path = os.path.join(get_user_folder(email), rel_path)
    if not os.path.exists(abs_path):
        return jsonify({'success': False, 'message': 'File/folder tidak ditemukan'}), 404
    info = {
        'nama': os.path.basename(abs_path),
        'path': rel_path,
        'tipe': 'folder' if os.path.isdir(abs_path) else 'file',
        'ukuran': os.path.getsize(abs_path) if os.path.isfile(abs_path) else None,
        'waktu_dibuat': datetime.fromtimestamp(os.path.getctime(abs_path)).strftime('%Y-%m-%d %H:%M'),
        'waktu_diubah': datetime.fromtimestamp(os.path.getmtime(abs_path)).strftime('%Y-%m-%d %H:%M'),
    }
    if os.path.isdir(abs_path):
        info['jumlah_item'] = len(os.listdir(abs_path))
    return jsonify({'success': True, 'properties': info})

@app.route('/save_word', methods=['POST'])
def save_word():
    if 'user' not in session:
        return jsonify({'success': False}), 401
    data = request.get_json()
    content = data.get('content', '')
    filename = secure_filename(data.get('filename', 'DokumenBaru.docx'))
    old_filename = secure_filename(data.get('old_filename', filename))
    user_folder = get_user_folder(session['user'])
    old_file_path = os.path.join(user_folder, old_filename)
    file_path = os.path.join(user_folder, filename)
    # Rename jika nama file berubah
    if filename != old_filename:
        base, ext = os.path.splitext(filename)
        i = 1
        new_file_path = file_path
        while os.path.exists(new_file_path):
            new_file_path = os.path.join(user_folder, f"{base}({i}){ext}")
            i += 1
        file_path = new_file_path
        if os.path.exists(old_file_path):
            os.rename(old_file_path, file_path)
    # Simpan konten baru (dengan gambar)
    doc = Document()
    soup = BeautifulSoup(content, 'html.parser')
    for el in soup.find_all(['p', 'h1', 'h2', 'h3', 'li']):
        para = doc.add_paragraph()
        for child in getattr(el, 'children', []):
            # Hanya proses jika child adalah Tag (bukan NavigableString)
            if hasattr(child, 'name') and child.name == 'img' and hasattr(child, 'has_attr') and child.has_attr('src'):
                src = child['src']
                if src.startswith('data:image/'):
                    # Simpan gambar base64 ke file sementara
                    match = re.match(r'data:image/(png|jpeg);base64,(.*)', src)
                    if match:
                        ext = match.group(1)
                        img_data = base64.b64decode(match.group(2))
                        img_folder = os.path.join(user_folder, 'word_images')
                        if not os.path.exists(img_folder):
                            os.makedirs(img_folder, exist_ok=True)
                        img_name = f"img_{hash(src)}.{ext}"
                        img_path = os.path.join(img_folder, img_name)
                        with open(img_path, 'wb') as f:
                            f.write(img_data)
                        para.add_run().add_picture(img_path, width=None)
                elif src.startswith('/word_image/'):
                    # Gambar sudah ada di server
                    img_path = os.path.join(user_folder, 'word_images', os.path.basename(src))
                    if os.path.exists(img_path):
                        para.add_run().add_picture(img_path, width=None)
            elif not hasattr(child, 'name'):
                para.add_run(str(child))
    doc.save(file_path)
    return jsonify({'success': True})

@app.route('/edit_word/<filename>')
def edit_word(filename):
    if 'user' not in session:
        return redirect('/login')
    user_folder = get_user_folder(session['user'])
    file_path = os.path.join(user_folder, filename)
    if not os.path.exists(file_path):
        flash('File tidak ditemukan')
        return redirect('/dashboard')
    from docx import Document
    from docx.opc.constants import RELATIONSHIP_TYPE as RT
    from base64 import b64encode
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    doc = Document(file_path)
    content = ''
    rels = doc.part.rels
    img_map = {}
    # Gambar
    for rel in rels.values():
        if rel.reltype == RT.IMAGE:
            img_bytes = rel.target_part.blob
            ext = rel.target_ref.split('.')[-1].lower()
            mime = 'image/png' if ext == 'png' else 'image/jpeg'
            data_url = f"data:{mime};base64,{b64encode(img_bytes).decode()}"
            img_map[rel.rId] = data_url
    # List detection helpers
    def is_bullet(para):
        return para.style.name.lower().startswith('list bullet')
    def is_number(para):
        return para.style.name.lower().startswith('list number')
    # Build HTML
    in_ul = in_ol = False
    for para in doc.paragraphs:
        html = ''
        # Alignment
        align = ''
        if para.alignment == WD_PARAGRAPH_ALIGNMENT.CENTER:
            align = ' style="text-align:center;"'
        elif para.alignment == WD_PARAGRAPH_ALIGNMENT.RIGHT:
            align = ' style="text-align:right;"'
        elif para.alignment == WD_PARAGRAPH_ALIGNMENT.JUSTIFY:
            align = ' style="text-align:justify;"'
        # Runs (formatting)
        for run in para.runs:
            run_html = run.text.replace('\n', '<br>')
            if not run_html:
                continue
            if run.bold:
                run_html = f'<b>{run_html}</b>'
            if run.italic:
                run_html = f'<i>{run_html}</i>'
            if run.underline:
                run_html = f'<u>{run_html}</u>'
            # Gambar inline
            for drawing in run._element.xpath('.//*[local-name()="blip"]'):
                rEmbed = drawing.attrib.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                if rEmbed and rEmbed in img_map:
                    run_html += f'<img src="{img_map[rEmbed]}" style="max-width:100%;max-height:300px;"/>'
            html += run_html
        style = para.style.name.lower()
        # Heading
        if 'heading' in style:
            level = ''.join(filter(str.isdigit, style)) or '1'
            content += f'<h{level}{align}>{html}</h{level}>'
        # List
        elif is_bullet(para):
            if not in_ul:
                content += '<ul>'
                in_ul = True
            content += f'<li{align}>{html}</li>'
        elif is_number(para):
            if not in_ol:
                content += '<ol>'
                in_ol = True
            content += f'<li{align}>{html}</li>'
        else:
            if in_ul:
                content += '</ul>'
                in_ul = False
            if in_ol:
                content += '</ol>'
                in_ol = False
            content += f'<p{align}>{html}</p>'
    if in_ul:
        content += '</ul>'
    if in_ol:
        content += '</ol>'
    return render_template('edit_word.html', content=content, filename=filename)

@app.route('/download_word/<filename>')
def download_word(filename):
    if 'user' not in session:
        return redirect('/login')
    user_folder = get_user_folder(session['user'])
    file_path = os.path.join(user_folder, filename)
    return send_file(file_path, as_attachment=True)

@app.route('/create_word', methods=['POST'])
def create_word():
    if 'user' not in session:
        return jsonify({'success': False, 'message': 'Not logged in'}), 401
    email = session['user']
    data = request.get_json() or {}
    rel_path = data.get('path', '').strip()
    if '..' in rel_path or rel_path.startswith('/'):
        return jsonify({'success': False, 'message': 'Path tidak valid'}), 400
    user_folder = get_user_folder(email)
    current_folder = os.path.join(user_folder, rel_path)
    if os.path.exists(current_folder):
        if not os.path.isdir(current_folder):
            return jsonify({'success': False, 'message': f'Path {current_folder} exists and is not a directory!'}), 400
    else:
        try:
            os.makedirs(current_folder, exist_ok=True)
        except FileExistsError:
            if not os.path.isdir(current_folder):
                return jsonify({'success': False, 'message': f'Path {current_folder} exists and is not a directory!'}), 400
    # Cari nama file unik
    base_name = 'DokumenBaru'
    ext = '.docx'
    i = 0
    while True:
        if i == 0:
            filename = base_name + ext
        else:
            filename = f"{base_name}({i}){ext}"
        file_path = os.path.join(current_folder, filename)
        if not os.path.exists(file_path):
            break
        i += 1
    # Buat dokumen Word kosong
    doc = Document()
    doc.add_paragraph('')
    doc.save(file_path)
    return jsonify({'success': True, 'filename': filename})

@app.route('/delete_file_word/<filename>', methods=['POST'])
def delete_file_word(filename):
    if 'user' not in session:
        return jsonify({'success': False, 'message': 'Not logged in'}), 401
    user_folder = get_user_folder(session['user'])
    file_path = os.path.join(user_folder, filename)
    if os.path.exists(file_path) and os.path.isfile(file_path):
        try:
            os.remove(file_path)
            return jsonify({'success': True})
        except Exception as e:
            return jsonify({'success': False, 'message': str(e)})
    return jsonify({'success': False, 'message': 'File tidak ditemukan'}), 404

@app.route('/upload_image_word', methods=['POST'])
def upload_image_word():
    if 'user' not in session:
        return jsonify({'error': 'Not logged in'}), 401
    if 'file' not in request.files:
        return jsonify({'error': 'No file uploaded'}), 400
    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'No file selected'}), 400
    filename = secure_filename(file.filename)
    user_folder = get_user_folder(session['user'])
    img_folder = os.path.join(user_folder, 'word_images')
    if not os.path.exists(img_folder):
        os.makedirs(img_folder, exist_ok=True)
    # Pastikan nama unik
    base, ext = os.path.splitext(filename)
    i = 1
    save_path = os.path.join(img_folder, filename)
    while os.path.exists(save_path):
        filename = f"{base}({i}){ext}"
        save_path = os.path.join(img_folder, filename)
        i += 1
    file.save(save_path)
    url = url_for('serve_word_image', filename=filename)
    return jsonify({'location': url})

@app.route('/word_image/<filename>')
def serve_word_image(filename):
    if 'user' not in session:
        return '', 403
    user_folder = get_user_folder(session['user'])
    img_folder = os.path.join(user_folder, 'word_images')
    return send_from_directory(img_folder, filename)

@app.route('/create_excel', methods=['POST'])
def create_excel():
    if 'user' not in session:
        return jsonify({'success': False, 'message': 'Not logged in'}), 401
    email = session['user']
    data = request.get_json() or {}
    rel_path = data.get('path', '').strip()
    if '..' in rel_path or rel_path.startswith('/'):
        return jsonify({'success': False, 'message': 'Path tidak valid'}), 400
    user_folder = get_user_folder(email)
    current_folder = os.path.join(user_folder, rel_path)
    if os.path.exists(current_folder):
        if not os.path.isdir(current_folder):
            return jsonify({'success': False, 'message': f'Path {current_folder} exists and is not a directory!'}), 400
    else:
        try:
            os.makedirs(current_folder, exist_ok=True)
        except FileExistsError:
            if not os.path.isdir(current_folder):
                return jsonify({'success': False, 'message': f'Path {current_folder} exists and is not a directory!'}), 400
    # Cari nama file unik
    base_name = 'ExcelBaru'
    ext = '.xlsx'
    i = 0
    while True:
        if i == 0:
            filename = base_name + ext
        else:
            filename = f"{base_name}({i}){ext}"
        file_path = os.path.join(current_folder, filename)
        if not os.path.exists(file_path):
            break
        i += 1
    # Buat file Excel kosong
    wb = Workbook()
    ws = wb.active
    ws.title = "Sheet1"
    wb.save(file_path)
    return jsonify({'success': True, 'filename': filename})

@app.route('/edit_excel/<filename>')
def edit_excel(filename):
    if 'user' not in session:
        return redirect('/login')
    user_folder = get_user_folder(session['user'])
    file_path = os.path.join(user_folder, filename)
    if not os.path.exists(file_path):
        flash('File tidak ditemukan')
        return redirect('/dashboard')
    from openpyxl import load_workbook
    wb = load_workbook(file_path)
    ws = wb.active
    # Ambil data sesuai jumlah baris dan kolom asli Excel
    max_row = ws.max_row
    max_col = ws.max_column
    data = []
    for row in ws.iter_rows(min_row=1, max_row=max_row, min_col=1, max_col=max_col, values_only=True):
        data.append(list(row))
    # Pad setiap baris agar jumlah kolom sama dengan max_col
    for row in data:
        if len(row) < max_col:
            row.extend([None] * (max_col - len(row)))
    # Tambahkan baris kosong jika data kurang dari 100 baris
    MIN_ROWS = 100
    while len(data) < MIN_ROWS:
        data.append([None] * max_col)
    return render_template('edit_excel.html', filename=filename, data=data)

@app.route('/delete_file_excel/<filename>', methods=['POST'])
def delete_file_excel(filename):
    if 'user' not in session:
        return jsonify({'success': False, 'message': 'Not logged in'}), 401
    user_folder = get_user_folder(session['user'])
    file_path = os.path.join(user_folder, filename)
    if os.path.exists(file_path) and os.path.isfile(file_path):
        try:
            os.remove(file_path)
            return jsonify({'success': True})
        except Exception as e:
            return jsonify({'success': False, 'message': str(e)})
    return jsonify({'success': False, 'message': 'File tidak ditemukan'}), 404

@app.route('/save_excel', methods=['POST'])
def save_excel():
    if 'user' not in session:
        return jsonify({'success': False}), 401
    data = request.get_json()
    filename = secure_filename(data.get('filename', 'ExcelBaru.xlsx'))
    old_filename = secure_filename(data.get('old_filename', filename))
    user_folder = get_user_folder(session['user'])
    old_file_path = os.path.join(user_folder, old_filename)
    file_path = os.path.join(user_folder, filename)
    # Rename jika nama file berubah
    if filename != old_filename:
        base, ext = os.path.splitext(filename)
        i = 1
        new_file_path = file_path
        while os.path.exists(new_file_path):
            new_file_path = os.path.join(user_folder, f"{base}({i}){ext}")
            i += 1
        file_path = new_file_path
        if os.path.exists(old_file_path):
            os.rename(old_file_path, file_path)
    # Simpan data ke file Excel
    from openpyxl import Workbook
    wb = Workbook()
    ws = wb.active
    ws.title = "Sheet1"
    luckysheet_data = data.get('data', [])
    for i, row in enumerate(luckysheet_data, 1):
        for j, cell in enumerate(row, 1):
            if cell and isinstance(cell, dict) and 'v' in cell:
                ws.cell(row=i, column=j, value=cell['v'])
    wb.save(file_path)
    return jsonify({'success': True})

@app.route('/devices')
def devices():
    if 'user' not in session:
        return redirect('/login')
    email = session['user']
    user = users.get(email)
    device_list = user.get('devices', []) if user else []
    # Tandai device aktif (session_token cocok dengan user['session_token'])
    for d in device_list:
        d['active'] = (d.get('session_token') == user.get('session_token'))
    return render_template('devices.html', devices=device_list)

@app.route('/settings')
def settings():
    if 'user' not in session:
        return redirect('/login')
    email = session['user']
    user = users.get(email)
    device_list = user.get('devices', []) if user else []
    for d in device_list:
        d['active'] = (d.get('session_token') == user.get('session_token'))
    return render_template('settings.html', devices=device_list)

@app.route('/payment', methods=['GET', 'POST'])
def payment():
    package = request.args.get('package', 'S')
    packages = {
        'S': {'name': 'Paket S', 'size': '50GB', 'price': 'Rp9.900'},
        'M': {'name': 'Paket M', 'size': '100GB', 'price': 'Rp14.900'},
        'L': {'name': 'Paket L', 'size': 'Coming Soon', 'price': '-'}
    }
    if package not in packages:
        package = 'S'
    if request.method == 'POST' and package in ['S', 'M', 'L']:
        session['user_has_package'] = True
        session['user_package_code'] = package
        flash('Pembelian paket berhasil!')
        return redirect(url_for('dashboard'))
    return render_template('payment.html', package=packages[package], code=package)

@app.route('/package/<code>')
def package_detail(code):
    packages = {
        'S': {
            'name': 'Paket S',
            'size': '50GB',
            'price': 'Rp9.900',
            'desc': 'Paket S cocok untuk pengguna personal yang ingin menyimpan dokumen, foto, dan file penting dengan kapasitas besar namun tetap hemat biaya.'
        },
        'M': {
            'name': 'Paket M',
            'size': '100GB',
            'price': 'Rp14.900',
            'desc': 'Paket M ideal untuk pengguna aktif, pelajar, atau UMKM yang membutuhkan ruang lebih untuk backup data, file kerja, dan media.'
        },
        'L': {
            'name': 'Paket L',
            'size': 'Coming Soon',
            'price': '-',
            'desc': 'Paket L akan hadir dengan kapasitas lebih besar dan fitur premium untuk kebutuhan bisnis dan tim.'
        }
    }
    if code not in packages:
        code = 'S'
    user_has_package = session.get('user_has_package', False)
    user_package_code = session.get('user_package_code', 'S' if not user_has_package else session.get('user_package_code', 'M'))
    return render_template('package_detail.html', package=packages[code], code=code, user_has_package=user_has_package, user_package_code=user_package_code)

@app.route('/delete_device_history', methods=['POST'])
def delete_device_history():
    if 'user' not in session:
        return redirect('/login')
    email = session['user']
    user = users.get(email)
    if not user:
        flash('User tidak ditemukan')
        return redirect('/settings')
    device_id = request.form.get('device_id')
    if device_id:
        # Hapus satu perangkat berdasarkan id
        devices = user.get('devices', [])
        new_devices = [d for d in devices if str(d.get('id')) != str(device_id)]
        user['devices'] = new_devices
        save_users()
        flash('Perangkat berhasil dihapus')
    else:
        # Hapus semua history perangkat
        user['devices'] = []
        save_users()
        flash('Seluruh history perangkat berhasil dihapus')
    return redirect('/settings')

def register_user(user_data):
    user = create_user(user_data)
    user.package = 'S'
    user.status = 'pro S'
    user.save()
    return user

if __name__ == '__main__':
    app.run(debug=True, host='0.0.0.0', port=5000)

